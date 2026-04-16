"""
app/routes/docs.py — Serve training and planning documents as rendered Markdown.

Converts .docx and .md files from the docs/ directory into structured JSON
so the frontend can display them in the Docs tab without any build step.

Routes:
  GET /docs/list         — list all available documents (slug, title, category)
  GET /docs/{slug}       — return a single document as { slug, title, sections[] }
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse

from app.security import require_local_request_dep
from fastapi import Depends

router = APIRouter(
    prefix="/docs",
    tags=["Docs"],
    dependencies=[Depends(require_local_request_dep)],
)

DOCS_DIR = Path(__file__).parent.parent.parent / "docs"
TRAINING_DOCX = Path(__file__).parent.parent.parent / "lucid-api-training.docx"


# ---------------------------------------------------------------------------
# Document registry — defines what gets exposed and in what order
# ---------------------------------------------------------------------------

_DOC_REGISTRY: list[dict[str, Any]] = [
    {
        "slug": "training",
        "title": "Lucid APIs: Training Documentation",
        "category": "Training",
        "path": TRAINING_DOCX,
        "description": "Three-part guide: Reference · Applied · Hands-On",
    },
    {
        "slug": "training-part3-revised",
        "title": "Part 3 (Revised): Hands-On Training",
        "category": "Training",
        "path": DOCS_DIR / "training-part3-revised.md",
        "description": "Corrected hands-on exercises accurate to the current app",
    },
    {
        "slug": "prd",
        "title": "Product Requirements Document",
        "category": "App Reference",
        "path": DOCS_DIR / "lucid-api-explorer-PRD-final.docx",
        "description": "What the app is, who it's for, and what it does",
    },
    {
        "slug": "backend",
        "title": "Backend Architecture",
        "category": "App Reference",
        "path": DOCS_DIR / "lucid-api-explorer-BACKEND-final.docx",
        "description": "Tech stack, auth flows, and API integration details",
    },
    {
        "slug": "troubleshooting",
        "title": "Troubleshooting Guide",
        "category": "Reference",
        "path": DOCS_DIR / "troubleshooting.md",
        "description": "Common errors and fixes",
    },
]


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def _docx_to_blocks(path: Path) -> list[dict[str, str]]:
    """
    Convert a .docx file to a flat list of content blocks.
    Each block has { type: 'heading1'|'heading2'|'heading3'|'paragraph'|'list_item'|'code', text: str }.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return [{"type": "paragraph", "text": "(python-docx not installed — cannot render .docx files)"}]

    doc = Document(str(path))
    blocks: list[dict[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        if "Heading 1" in style:
            block_type = "heading1"
        elif "Heading 2" in style:
            block_type = "heading2"
        elif "Heading 3" in style:
            block_type = "heading3"
        elif "List" in style:
            block_type = "list_item"
        elif style in ("Code", "Preformatted Text"):
            block_type = "code"
        else:
            block_type = "paragraph"
        blocks.append({"type": block_type, "text": text})

    # Also pull table data as paragraph rows
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append({"type": "table", "text": "\n".join(rows)})

    return blocks


def _md_to_blocks(path: Path) -> list[dict[str, str]]:
    """
    Parse a Markdown file into content blocks.
    Handles headings, fenced code blocks, list items, and paragraphs.
    """
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    blocks: list[dict[str, str]] = []
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if in_code_block:
                blocks.append({"type": "code", "text": "\n".join(code_lines)})
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("#### "):
            blocks.append({"type": "heading3", "text": stripped[5:]})
        elif stripped.startswith("### "):
            blocks.append({"type": "heading3", "text": stripped[4:]})
        elif stripped.startswith("## "):
            blocks.append({"type": "heading2", "text": stripped[3:]})
        elif stripped.startswith("# "):
            blocks.append({"type": "heading1", "text": stripped[2:]})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append({"type": "list_item", "text": stripped[2:]})
        elif re.match(r"^\d+\. ", stripped):
            blocks.append({"type": "list_item", "text": re.sub(r"^\d+\. ", "", stripped)})
        elif re.match(r"^\|.+\|", stripped):
            # Markdown table row — collect the whole table
            table_rows = [stripped]
            while i + 1 < len(lines) and re.match(r"^\|.+\|", lines[i + 1].strip()):
                i += 1
                row = lines[i].strip()
                # Skip separator rows like |---|---|
                if not re.match(r"^\|[-| :]+\|$", row):
                    table_rows.append(row)
            # Strip leading/trailing pipes and normalize cells
            cleaned = []
            for row in table_rows:
                cells = [c.strip() for c in row.strip("|").split("|")]
                if any(cells):
                    cleaned.append(" | ".join(cells))
            if cleaned:
                blocks.append({"type": "table", "text": "\n".join(cleaned)})
        else:
            blocks.append({"type": "paragraph", "text": stripped})

        i += 1

    return blocks


def _load_doc(entry: dict[str, Any]) -> dict[str, Any]:
    """Load and convert a document entry to its block representation."""
    path: Path = entry["path"]
    if not path.exists():
        blocks = [{"type": "paragraph", "text": f"File not found: {path.name}"}]
    elif path.suffix == ".docx":
        blocks = _docx_to_blocks(path)
    elif path.suffix == ".md":
        blocks = _md_to_blocks(path)
    else:
        blocks = [{"type": "paragraph", "text": path.read_text(encoding="utf-8")}]

    return {
        "slug": entry["slug"],
        "title": entry["title"],
        "category": entry["category"],
        "description": entry["description"],
        "blocks": blocks,
    }


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.get("/list")
async def list_docs() -> JSONResponse:
    """Return metadata for all available documents (no content)."""
    docs = [
        {
            "slug": e["slug"],
            "title": e["title"],
            "category": e["category"],
            "description": e["description"],
            "available": e["path"].exists(),
        }
        for e in _DOC_REGISTRY
    ]
    return JSONResponse({"ok": True, "docs": docs})


@router.get("/{slug}")
async def get_doc(slug: str) -> JSONResponse:
    """Return a single document as structured blocks."""
    entry = next((e for e in _DOC_REGISTRY if e["slug"] == slug), None)
    if entry is None:
        raise HTTPException(status_code=404, detail=f"Document '{slug}' not found")
    doc = _load_doc(entry)
    return JSONResponse({"ok": True, "doc": doc})
